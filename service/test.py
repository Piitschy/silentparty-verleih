from docx import Document
import re

S_REPLACE:list[str] = ["{{","}}"]
S_FORMULA:list[str] = ["{%","%}"]
S_TABLE:str = S_FORMULA[0]+'TABLE'+S_FORMULA[1]

DEMO:dict = {
  "organisation": "Orga",
  "vertreten_durch": "Vertreter",
  "strasse": "Musterstr",
  "hausnummer":"1",
  "plz": "30167",
  "ort": "Hannover",
  "d_angebot": "01.07.2022",
  "d_guetlig_bis": "02.07.2022",
  "d_nutzung_start": "03.07.2022",
  "d_nutzung_end": "04.07.2022",
  "p_reinigung": 1,
  "p_headphone": 25,
  "p_kabel": 5,
  "p_sender": 200,
  "order": [
    {
      "pos":1,
      "text":"Kopfhörer",
      "subtext": "",
      "menge": 100,
      "p_einzel": 1.8,
      "summe": 180
    },
    {
      "pos":2,
      "text":"Sender",
      "subtext": "",
      "menge": 3,
      "p_einzel": 30,
      "summe": 90
    },
    {
      "pos":3,
      "text":"Aufwandspauschale",
      "subtext": "(Laden und Reinigen der KH nach der Vermietung, Ausgabe und Annahme der vermieteten Gegenstände)",
      "menge": 1,
      "p_einzel": 50,
      "summe": 50
    }
  ]
}

def btwn(text,sep:list[str]=S_REPLACE):
  return f'{sep[0]}{text}{sep[1]}'

def replace_text(text:str, backbone:dict=DEMO, parentKey:str=None, enumeration:int=None) -> str:
  if S_REPLACE[0] not in text and S_REPLACE[1] not in text:
    return text
  for key, value in backbone.items():
    if key in text:
      text_types = [str,int]
      if type(value) in text_types:
        text = text.replace(btwn(key), str(value))
        if parentKey:
          if enumeration:
            text = text.replace(btwn(f'{parentKey}[{enumeration}].{key}'), str(value))
          else:
            text = text.replace(btwn(f'{parentKey}.{key}'), str(value))
      elif type(value) == list:
        for i,e in enumerate(value):
          text = replace_text(text, e, parentKey=key, enumeration=i+1)
      elif type(value) == dict:
        text = replace_text(text, e, parentKey=key)
  return text

def replace_in_paragraph(p) -> str:
  inline = p.runs
  for i in range(len(inline)):
    text = replace_text(inline[i].text)
    inline[i].text = text
  return p

doc = Document('./templates/angebot.docx')

for table in doc.tables:
  first_cell = table.rows[0].cells[0]
  text = first_cell.text
  key, value = None
  params = []
  for k, v in DEMO.items():
    if S_REPLACE[0]+key in text:
      key, value = k, v
      break
  if '|' in text:
    start: int = text.find('|')
    end: int = text.find(S_REPLACE[1])
    param_s = text[start:end].replace(' ','')
    params = param_s.split(',')
  
  row_number = 0 if 'headline' not in params else 0
  table.rows[row_number]
  for e in DEMO[key]:
    pass
  

for table in doc.tables:
  for row in table.rows:
    for cell in row.cells:
      for p in cell.paragraphs:
        p = replace_in_paragraph(p)

for p in doc.paragraphs:
  p = replace_in_paragraph(p)
  

doc.save('./temporary/test.docx')